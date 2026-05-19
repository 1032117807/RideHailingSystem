# -*- coding: utf-8 -*-
import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
RESULTS = ROOT / "loadtests" / "results"
OUT = RESULTS / f"TripVerse-k6-loadtest-report-8081-fixed-{datetime.now().strftime('%Y%m%d-%H%M%S')}.docx"

RUN_FILES = [
    ("10 VU 冒烟", "tripverse-core-10vus-8081-summary.json"),
    ("200 VU", "tripverse-core-200vus-8081-summary.json"),
    ("500 VU", "tripverse-core-500vus-8081-summary.json"),
    ("1000 VU", "tripverse-core-1000vus-8081-summary.json"),
    ("1200 VU", "tripverse-core-1200vus-8081-summary.json"),
]

ENDPOINT_NAMES = {
    "endpoint_driver_income": "司机收入",
    "endpoint_driver_dashboard": "司机工作台",
    "endpoint_admin_dashboard": "管理员仪表盘",
    "endpoint_admin_users_summary": "用户汇总",
    "endpoint_admin_risk_logs": "风控日志",
    "endpoint_passenger_orders_my": "乘客订单",
    "endpoint_admin_orders": "管理员订单",
    "endpoint_admin_tokens": "Token 用量",
    "endpoint_public_search_main": "主线路票务搜索",
    "endpoint_public_search_transfer": "换乘票务搜索",
    "endpoint_public_ticket_detail": "票务详情",
    "endpoint_driver_trips": "司机行程列表",
    "endpoint_admin_users": "管理员用户列表",
}


def fmt_ms(value):
    return "-" if value is None else f"{float(value):,.2f}"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def write_cell(cell, value, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(value))
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def load_runs():
    runs = []
    for label, filename in RUN_FILES:
        metrics = json.loads((RESULTS / filename).read_text(encoding="utf-8"))["metrics"]
        duration = metrics["http_req_duration"]
        fail_rate = float(metrics["http_req_failed"]["value"])
        runs.append(
            {
                "label": label,
                "file": filename,
                "vus": int(metrics["vus_max"]["value"]),
                "requests": int(metrics["http_reqs"]["count"]),
                "rps": float(metrics["http_reqs"]["rate"]),
                "avg": duration.get("avg"),
                "p90": duration.get("p(90)"),
                "p95": duration.get("p(95)"),
                "max": duration.get("max"),
                "failed": fail_rate,
                "checks": float(metrics["checks"]["value"]),
                "passed_latency": duration.get("p(95)", 0) < 1000,
                "passed_failed": fail_rate < 0.05,
                "metrics": metrics,
            }
        )
    return runs


def main():
    runs = load_runs()
    endpoint_keys = [
        key
        for key, value in runs[-1]["metrics"].items()
        if key.startswith("endpoint_") and isinstance(value, dict) and "p(95)" in value
    ]
    worst_endpoints = sorted(
        endpoint_keys, key=lambda key: runs[-1]["metrics"][key].get("p(95)", 0), reverse=True
    )[:8]

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    for style_name in ["Normal", "Title", "Heading 1", "Heading 2"]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    doc.styles["Normal"].font.size = Pt(10)

    title = doc.add_heading("TripVerse 后端 k6 压测评估报告", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(
        "目标服务：http://127.0.0.1:8081/api    测试工具：k6 v2.0.0    测试日期：2026-05-14"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("一、结论摘要", level=1)
    add_bullets(
        doc,
        [
            "本轮按 10、200、500、1000、1200 VU 阶梯逐渐增加并发，核心读接口压测均已完成。",
            "10 VU 冒烟稳定，p95 为 72.01 ms，错误率 0%。",
            "从 200 VU 开始，p95 达到 2.03 s，已经超过脚本设定的 p95 < 1s 性能阈值。",
            "500 VU 后吞吐进入平台期，约 285-298 req/s；继续升到 1000/1200 VU，吞吐提升有限，但延迟和 EOF 失败明显增加。",
            "1200 VU 下错误率 2.00%，低于 5% 可用性阈值，但 p95 为 8.79 s，最大响应约 25.85 s，不建议作为可接受生产容量。",
            "按“p95 < 1s 且错误率 < 5%”口径，本环境的可接受并发上限低于 200 VU；按“服务不崩且错误率 < 5%”口径，可顶到约 1200 VU，但用户体验较差。",
        ],
    )

    doc.add_heading("二、测试模型", level=1)
    add_bullets(
        doc,
        [
            "测试脚本：loadtests/tripverse-core.js。",
            "场景模型：ramping-vus 阶梯升压，每轮包含 ramp up、hold、ramp down。",
            "接口覆盖：公开票务查询、票务详情、乘客个人与订单、司机行程/工作台/收入、管理员仪表盘/用户/订单/Token/风控/知识库等读接口。",
            "每次业务迭代约 19 个 HTTP 请求，并包含 1 秒 sleep；setup 阶段执行乘客、司机、管理员登录并获取首个票务 ID。",
            "阈值：HTTP 失败率 < 5%，HTTP p95 < 1000 ms。",
            "后端当前代码中 database/sql 连接池最大打开连接数为 20，这对高并发读接口有明显约束。",
        ],
    )

    doc.add_heading("三、阶梯压测结果", level=1)
    table = doc.add_table(rows=1, cols=10)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["阶段", "峰值 VU", "请求数", "吞吐 req/s", "平均 ms", "p90 ms", "p95 ms", "最大 ms", "失败率", "阈值结论"]
    for index, header in enumerate(headers):
        write_cell(table.rows[0].cells[index], header, True)
        shade(table.rows[0].cells[index], "D9EAF7")

    for run in runs:
        row = table.add_row().cells
        threshold = "通过" if run["passed_latency"] and run["passed_failed"] else "延迟未达标"
        values = [
            run["label"],
            run["vus"],
            run["requests"],
            f"{run['rps']:.2f}",
            fmt_ms(run["avg"]),
            fmt_ms(run["p90"]),
            fmt_ms(run["p95"]),
            fmt_ms(run["max"]),
            f"{run['failed'] * 100:.2f}%",
            threshold,
        ]
        for index, value in enumerate(values):
            write_cell(row[index], value)
        if not run["passed_latency"]:
            shade(row[6], "FCE4D6")
        if run["failed"] > 0:
            shade(row[8], "FFF2CC")

    doc.add_heading("四、容量趋势判断", level=1)
    add_bullets(
        doc,
        [
            "吞吐趋势：200 VU 为 248.19 req/s，500 VU 为 284.57 req/s，1000 VU 为 291.39 req/s，1200 VU 为 297.96 req/s。500 VU 以后吞吐提升很小，说明后端或数据库已经接近瓶颈。",
            "延迟趋势：p95 从 200 VU 的 2.03 s 升至 500 VU 的 4.82 s，再升至 1000/1200 VU 的 8.52/8.79 s。并发越高，请求排队越明显。",
            "可用性趋势：1000 VU 开始出现 EOF，1200 VU 错误率为 2.00%，虽然仍低于 5%，但已经暴露连接或超时层面的稳定性风险。",
        ],
    )

    doc.add_heading("五、瓶颈接口", level=1)
    doc.add_paragraph("1200 VU 阶段 p95 最慢接口如下：")
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["指标名", "业务含义", "平均 ms", "p90 ms", "p95 ms", "最大 ms"]
    for index, header in enumerate(headers):
        write_cell(table.rows[0].cells[index], header, True)
        shade(table.rows[0].cells[index], "D9EAD3")
    for key in worst_endpoints:
        metric = runs[-1]["metrics"][key]
        row = table.add_row().cells
        values = [
            key,
            ENDPOINT_NAMES.get(key, "-"),
            fmt_ms(metric.get("avg")),
            fmt_ms(metric.get("p(90)")),
            fmt_ms(metric.get("p(95)")),
            fmt_ms(metric.get("max")),
        ]
        for index, value in enumerate(values):
            write_cell(row[index], value)
        if metric.get("p(95)", 0) > 5000:
            shade(row[4], "FCE4D6")

    doc.add_paragraph(
        "1000 VU 阶段同样主要集中在司机工作台、司机收入、管理员仪表盘、用户汇总等聚合类接口，说明高并发下慢点不是单个异常接口，而是聚合查询、数据库连接池和服务端排队共同作用。"
    )

    doc.add_heading("六、问题评估", level=1)
    add_bullets(
        doc,
        [
            "数据库连接池偏小：main.go 中 MaxOpenConns 为 20，高并发下大量请求等待数据库连接，表现为吞吐平台期和聚合接口长尾延迟。",
            "聚合查询接口成本较高：司机收入、司机工作台、管理员仪表盘、用户汇总等接口 p95 达到 8-25 s，优先排查这些接口的 SQL、索引和统计逻辑。",
            "服务端超时/连接稳定性风险：1000 VU 起出现 EOF，1200 VU 错误率为 2.00%，建议结合服务日志定位是否为请求超时、连接被关闭或数据库等待过长。",
            "当前脚本是混合读场景，不能代表下单、支付等写入场景容量；写入场景通常会更敏感，需要单独压测。",
        ],
    )

    doc.add_heading("七、优化建议", level=1)
    add_bullets(
        doc,
        [
            "先做数据库侧优化：为高频筛选、关联和排序字段补索引；对 dashboard、income、summary 类接口查看 EXPLAIN，减少全表扫描和重复聚合。",
            "调整连接池并复测：在数据库可承受的前提下，将 MaxOpenConns 从 20 分档提升到 50/100，配合 MaxIdleConns 调整，观察吞吐和 p95 是否改善。",
            "给聚合数据加缓存或预计算：管理员仪表盘、司机收入、用户汇总、Token 用量这类读多写少指标适合短 TTL 缓存或后台预聚合。",
            "增加接口级压测：把公开查询、司机工作台、管理员仪表盘拆成单接口压测，明确每个接口的独立容量。",
            "补充服务端观测：压测时同步采集 CPU、内存、MySQL 慢查询、连接池等待、goroutine 数和错误日志，才能精确区分应用瓶颈与数据库瓶颈。",
        ],
    )

    doc.add_heading("八、原始结果文件", level=1)
    for run in runs:
        doc.add_paragraph(f"{run['label']}: loadtests/results/{run['file']}", style="List Bullet")

    doc.add_heading("附录：执行命令", level=1)
    command = (
        "$env:BASE_URL='http://127.0.0.1:8081/api'; "
        "$env:PEAK_VUS='<VU>'; "
        "$env:RAMP_UP='<ramp>'; "
        "$env:HOLD='<hold>'; "
        "$env:RAMP_DOWN='<down>'; "
        "& 'C:\\Program Files\\k6\\k6.exe' run --quiet "
        "--summary-export 'loadtests\\results\\<name>.json' 'loadtests\\tripverse-core.js'"
    )
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(command)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(8)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
